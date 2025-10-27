import google.generativeai as genai
import requests
import pdfplumber
from docx import Document
from docx.shared import Pt, Inches # Added Inches for margins
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
import os
import time
import json
from typing import TypedDict, List, Dict, Optional # For state definition
import arxiv # Make sure arxiv is imported

# Import LangGraph components
from langgraph.graph import StateGraph, END

# --- Configuration ---
GEMINI_API_KEY = "" # <-- PUT YOUR GEMINI API KEY HERE
SEMANTIC_SCHOLAR_API_KEY = "" # <-- PUT YOUR SEMANTIC SCHOLAR KEY HERE
GEMINI_MODEL_NAME = "models/gemini-2.5-flash"
SEMANTIC_SCHOLAR_API_URL = "https://api.semanticscholar.org/graph/v1/paper/search"
PDF_DOWNLOAD_DIR = "downloaded_pdfs" # Directory to save PDFs temporarily
OUTPUT_REPORT_NAME = "AI_Generated_Research_Report_Reverted_Enhanced.docx" # New filename

# --- State Definition (Includes topic_overview) ---
class ResearchState(TypedDict):
    topic: str
    paper_limit: Optional[int]
    papers_to_process: Optional[List[Dict]]
    parsed_papers: Optional[List[Dict]]
    all_analyses: Optional[List[Dict]]
    topic_overview: Optional[str] # Added for overview
    comparison_result: Optional[str]
    report_path: Optional[str]
    error: Optional[str]

# --- Agent Functions ---

# --- Search Agent (Includes arXiv fallback - Keep this version) ---
def search_academic_papers_node(state: Dict):
    """
    Node function for the Search Agent. Queries Semantic Scholar and arXiv.
    Prioritizes papers with direct PDF links.
    """
    print("\n--- Node: Search Agent ---")
    topic = state.get("topic")
    limit = state.get("paper_limit", 3) # Target number of papers with PDFs
    papers_found_details = [] # Stores {paperId, title, pdf_url}

    if not topic:
        print("ERROR: Topic not found in state.")
        return {"papers_to_process": [], "error": "Topic missing"} # Ensure return has expected keys

    print(f"Searching for topic: '{topic}' (Targeting {limit} papers with PDFs)")

    # --- 1. Query Semantic Scholar ---
    try:
        print("Querying Semantic Scholar API...")
        ss_params = {'query': topic, 'limit': limit * 2, 'fields': 'title,paperId,externalIds,openAccessPdf'} # Get more results initially, ask for arXiv ID
        ss_headers = {'Accept': 'application/json', 'x-api-key': SEMANTIC_SCHOLAR_API_KEY}

        if SEMANTIC_SCHOLAR_API_KEY == "YOUR_SEMANTIC_SCHOLAR_API_KEY":
             raise ValueError("Semantic Scholar API Key not set.")

        time.sleep(1.1) # Respect rate limit
        ss_response = requests.get(SEMANTIC_SCHOLAR_API_URL, params=ss_params, headers=ss_headers)
        ss_response.raise_for_status()
        ss_results = ss_response.json()

        if ss_results and 'data' in ss_results and ss_results['data']:
            print(f"Semantic Scholar found {len(ss_results['data'])} potential papers.")
            for paper in ss_results['data']:
                paper_id = paper.get('paperId')
                title = paper.get('title')
                pdf_url = None

                # Check for direct Open Access PDF link first
                if paper.get('openAccessPdf') and paper['openAccessPdf'].get('url'):
                    pdf_url = paper['openAccessPdf']['url']
                    print(f"  Found Open Access PDF via S2: '{title}'")
                # Check if it has an arXiv ID listed in externalIds
                elif paper.get('externalIds') and paper['externalIds'].get('ArXiv'):
                    arxiv_id = paper['externalIds']['ArXiv']
                    pdf_url = f"https://arxiv.org/pdf/{arxiv_id}.pdf"
                    print(f"  Found arXiv ID via S2, constructing PDF link: '{title}' ({arxiv_id})")

                if paper_id and title and pdf_url and len(papers_found_details) < limit:
                    if not any(p.get('paperId') == paper_id for p in papers_found_details if p.get('paperId')):
                         papers_found_details.append({
                            'paperId': paper_id,
                            'title': title,
                            'pdf_url': pdf_url
                         })

            print(f"Collected {len(papers_found_details)} papers with PDF links via Semantic Scholar.")

    except Exception as e:
        print(f"Warning: Semantic Scholar search failed or partially failed: {type(e).__name__} - {e}")

    # --- 2. Query arXiv (if needed) ---
    if len(papers_found_details) < limit:
        needed = limit - len(papers_found_details)
        print(f"\nAttempting to find {needed} more papers via arXiv API...")
        try:
            arxiv_client = arxiv.Client()
            search = arxiv.Search(
                query=topic,
                max_results=needed * 2, # Get extra
                sort_by=arxiv.SortCriterion.Relevance
            )
            arxiv_results_count = 0
            for result in arxiv_client.results(search):
                if len(papers_found_details) >= limit: break

                arxiv_id = result.entry_id.split('/')[-1]
                title = result.title
                pdf_url = result.pdf_url
                is_duplicate = any(p['title'].lower()[:50] == title.lower()[:50] for p in papers_found_details)

                if pdf_url and not is_duplicate:
                    print(f"  Found paper via arXiv: '{title}'")
                    papers_found_details.append({
                        'paperId': f"arXiv:{arxiv_id}",
                        'title': title,
                        'pdf_url': pdf_url
                    })
                    arxiv_results_count += 1
            print(f"Collected {arxiv_results_count} additional papers via arXiv.")
        except Exception as e:
            print(f"Warning: arXiv search failed: {type(e).__name__} - {e}")

    # --- Final Check ---
    if not papers_found_details:
         print("ERROR: No papers with PDF links found from any source.")
         return {"papers_to_process": [], "error": "No papers with PDF links found"}

    print(f"\nTotal papers found with PDF links: {len(papers_found_details)}")
    return {"papers_to_process": papers_found_details, "error": None}

# --- PDF Downloader & Parser Agent (Reverted to working version) ---
def download_and_parse_pdfs_node(state: Dict):
    """Node function to download and parse PDFs (using the structure that previously worked)."""
    print("\n--- Node: Download & Parse Agent ---")
    papers_to_process = state.get("papers_to_process", [])
    parsed_papers = [] # List to store {paperId, title, text}
    final_error = state.get("error") # Preserve previous error state

    if not papers_to_process:
        if final_error:
             print("Skipping Download/Parse due to previous error.")
             return {"parsed_papers": [], "error": final_error}
        else:
             print("No papers to process found in state.")
             return {"parsed_papers": [], "error": "Search Agent found no processable papers."}

    # Create download directory if it doesn't exist
    os.makedirs(PDF_DOWNLOAD_DIR, exist_ok=True)
    print(f"Ensured download directory exists: '{PDF_DOWNLOAD_DIR}'")

    download_or_parse_failed_count = 0
    for paper in papers_to_process:
        paper_id = paper.get('paperId', 'unknown_id')
        pdf_url = paper.get('pdf_url')
        title = paper.get('title', 'Unknown Title')
        # Create a safe filename
        safe_title = "".join(c if c.isalnum() else "_" for c in title[:50])
        pdf_filename = f"{safe_title}_{paper_id[:8].replace(':', '_')}.pdf" # Ensure ':' from arXiv ID is replaced
        pdf_filepath = os.path.join(PDF_DOWNLOAD_DIR, pdf_filename)
        paper_text = "" # Initialize paper_text inside the loop for each paper

        if not pdf_url:
            print(f"Skipping paper '{title}' ({paper_id}) - Missing PDF URL.")
            download_or_parse_failed_count += 1
            continue

        print(f"\nProcessing paper: '{title}' ({paper_id})")

        try:
            # --- Download PDF ---
            print(f"Downloading PDF from: {pdf_url}")
            time.sleep(0.5) # Small delay before download
            pdf_response = requests.get(pdf_url, stream=True, timeout=90) # Stream for large files, 90s timeout
            pdf_response.raise_for_status()

            with open(pdf_filepath, 'wb') as f:
                for chunk in pdf_response.iter_content(chunk_size=8192):
                    f.write(chunk)
            print(f"Saved PDF to: '{pdf_filepath}'")

            # --- Parse PDF ---
            print(f"Parsing PDF: '{pdf_filepath}'")
            # Initialize text for THIS paper here
            current_paper_full_text = ""
            with pdfplumber.open(pdf_filepath) as pdf:
                print(f"  Opened PDF. Pages: {len(pdf.pages)}") # Debug print
                for i, page in enumerate(pdf.pages):
                    # Extract text for the current page
                    page_text = page.extract_text(x_tolerance=1, y_tolerance=1) # Extract text for *this* page
                    if page_text: # Check if text was extracted for *this* page
                        current_paper_full_text += page_text + "\n" # Append page text
                    # else:
                    #    print(f"  Warning: No text extracted from page {i+1}.")
            paper_text = current_paper_full_text # Assign accumulated text

            if paper_text:
                print(f"Successfully parsed. Text length: {len(paper_text)} chars.")
                parsed_papers.append({
                    "paperId": paper_id,
                    "title": title,
                    "text": paper_text
                })
            else:
                print(f"Warning: No text extracted from PDF: '{pdf_filepath}'")
                download_or_parse_failed_count += 1

            # --- Clean up downloaded file ---
            try:
                os.remove(pdf_filepath)
                print(f"Deleted temporary PDF: '{pdf_filepath}'")
            except OSError as e:
                print(f"Warning: Could not delete temporary PDF '{pdf_filepath}': {e}")

        except Exception as e:
            print(f"ERROR processing paper {paper_id} ('{title}'): {type(e).__name__} - {e}")
            download_or_parse_failed_count += 1
            # Clean up failed download/parse attempt if file exists
            if os.path.exists(pdf_filepath):
                try: os.remove(pdf_filepath); print(f"Cleaned up failed file: '{pdf_filepath}'")
                except OSError: pass

    print(f"\nFinished processing. Papers successfully parsed: {len(parsed_papers)}. Failures: {download_or_parse_failed_count}")
    if len(parsed_papers) == 0 and len(papers_to_process) > 0 and not final_error: # Set error only if none existed before
        final_error = "Failed to download or parse any valid PDFs."

    return {"parsed_papers": parsed_papers, "error": final_error}

# --- Summarizer Agent (Keep previous plain text version) ---
def summarize_papers_node(state: Dict):
    """Node function for the Summarizer Agent."""
    print("\n--- Node: Summarizer Agent ---")
    parsed_papers = state.get("parsed_papers", [])
    all_analyses = []
    final_error = state.get("error")

    if not parsed_papers:
        if final_error: return {"all_analyses": [], "error": final_error}
        else: return {"all_analyses": [], "error": "Parsing step yielded no text."}

    print(f"Summarizing {len(parsed_papers)} papers...")
    if GEMINI_API_KEY == "YOUR_GEMINI_API_KEY": return {"error": "Gemini API Key missing"}
    try:
        genai.configure(api_key=GEMINI_API_KEY)
        model = genai.GenerativeModel(GEMINI_MODEL_NAME)
    except Exception as e: return {"error": f"Gemini configuration failed: {e}"}

    summarization_failures = 0
    for i, paper in enumerate(parsed_papers):
        print(f"\nSummarizing paper {i+1}/{len(parsed_papers)}: '{paper.get('title', 'N/A')}'")
        paper_text = paper.get("text", "")
        if not paper_text: print("Skipping - paper has no text."); summarization_failures += 1; continue
        try:
            # --- Create Prompt (Plain Text Output) ---
            prompt = f"""Analyze the research paper text provided below. Extract the following information and present it clearly under the specified headings.

**Summary:**
Provide a concise summary (3-5 sentences) covering the paper's main objectives, methods, and key conclusions.

**Methodology:**
Briefly describe the core methodology, algorithms, or experimental approach used.

**Key Findings:**
List the 2-3 most significant findings, results, or outcomes reported in the paper.

--- START OF PAPER TEXT ---
{paper_text[:40000]}
--- END OF PAPER TEXT ---

Provide the output clearly structured under the headings "Summary:", "Methodology:", and "Key Findings:". Ensure each section is clearly separated. Do NOT use markdown formatting like asterisks for bolding in your response text.
"""
            print("Sending request to Gemini API...")
            response = model.generate_content(prompt, request_options={"timeout": 180})
            if response.parts:
                generated_text = response.text
                print("Received response from Gemini.")
                # Basic parsing
                summary = generated_text.split("Methodology:")[0].replace("Summary:", "").strip()
                methodology = generated_text.split("Key Findings:")[0].split("Methodology:")[-1].strip()
                key_findings = generated_text.split("Key Findings:")[-1].strip()
                all_analyses.append({'title': paper.get('title', 'N/A'), 'pdf_url': paper.get('pdf_url', 'N/A'), 'summary': summary, 'methodology': methodology, 'key_findings': key_findings})
            elif response.prompt_feedback:
                print(f"Gemini request blocked: {response.prompt_feedback.block_reason}")
                summarization_failures += 1
            else:
                print(f"Received empty response from Gemini.")
                summarization_failures += 1
        except Exception as e:
            print(f"ERROR summarizing paper '{paper.get('title', 'N/A')}': {e}")
            summarization_failures += 1

    print(f"\nFinished summarizing. Analyses generated: {len(all_analyses)}. Failures: {summarization_failures}")
    if len(all_analyses) == 0 and len(parsed_papers) > 0 and not final_error:
        final_error = "Summarization failed for all parsed papers."
    return {"all_analyses": all_analyses, "error": final_error}

# --- Comparator Agent (ENHANCED Prompt for Overview and Comparison - Plain Text Output) ---
def compare_analyses_node(state: Dict):
    """
    Node function for the Comparator Agent. Generates both a Topic Overview
    and a Detailed Comparison as plain text sections.
    """
    print("\n--- Node: Comparator Agent ---")
    all_analyses = state.get("all_analyses", [])
    topic = state.get("topic", "the research field")
    topic_overview_result = "Overview could not be generated." # Default
    comparison_result = "Detailed comparison could not be generated." # Default
    final_error = state.get("error")

    if not all_analyses:
        if final_error: return {"topic_overview": "", "comparison_result": "", "error": final_error}
        else: return {"topic_overview": "N/A - No papers analyzed.", "comparison_result": "N/A - No papers analyzed.", "error": "No summaries available"}

    # Handle single paper case
    if len(all_analyses) == 1:
        print("Only one paper analyzed. Generating overview only.")
        comparison_result = "Comparison skipped: Only one paper was successfully analyzed."
    else:
        print(f"Generating Overview & Comparison for {len(all_analyses)} papers...")

    if GEMINI_API_KEY == "YOUR_GEMINI_API_KEY": return {"error": "Gemini API Key missing"}

    try:
        genai.configure(api_key=GEMINI_API_KEY)
        model = genai.GenerativeModel(GEMINI_MODEL_NAME)

        # --- Prepare Input Text ---
        analysis_input = "Analysis details from the retrieved papers:\n\n"
        for i, analysis in enumerate(all_analyses):
            analysis_input += f"--- Paper {i+1} ({analysis.get('title', 'N/A')}) ---\n"
            analysis_input += f"Summary: {analysis.get('summary', 'N/A')}\n"
            analysis_input += f"Methodology: {analysis.get('methodology', 'N/A')}\n"
            analysis_input += f"Key Findings: {analysis.get('key_findings', 'N/A')}\n\n"

        # --- Create ENHANCED Prompt for BOTH Overview and Comparison (Plain Text Output) ---
        # --- (Inside compare_analyses_node function, within the try block) ---
        # --- Create ENHANCED Prompt for Overview and Comparison (Encourage Points) ---
        prompt = f"""Based *only* on the provided analysis details from the retrieved papers below, generate two distinct sections: "Topic Overview" and "Detailed Comparative Analysis".

{analysis_input}
--- Analysis Tasks ---

**1. Topic Overview:**
Synthesize insights from ALL provided papers to give a high-level overview of the research topic '{topic}'. Use bullet points for lists where appropriate (e.g., for methods, findings, future work). Address:
    * Current State/Focus: General status or recent focus.
    * Methods Used: Common techniques, algorithms, frameworks.
    * Performance/Accuracy: Key quantitative results or KPIs mentioned.
    * Accomplishments/Key Findings: Significant findings highlighted collectively.
    * Future Work/Directions: Common limitations or future research suggestions.

**2. Detailed Comparative Analysis:** (Only generate meaningful content here if more than one paper was provided)
Provide a point-by-point comparison *between* the papers (if more than one). Use bullet points within each comparison category where appropriate. Analyze and highlight:
    * Core Objective & Scope: Compare goals and focus.
    * Methodology & Approach: Compare specific methods, novelty, complexity, tools.
    * Key Findings & Performance: Compare quantitative results and qualitative findings.
    * Advancements & Relation to State-of-the-Art: Compare novelty or improvements.
    * Limitations & Future Work: Compare stated limitations or future directions.
    * Overall Contribution & Theme: Compare their main contributions.

Structure your entire response clearly with the exact headings "Topic Overview:" followed by its content, and then "Detailed Comparative Analysis:" followed by its content. Do NOT use markdown formatting like asterisks for bolding within the generated text content itself. Ensure lists naturally use bullet points or numbered formats.
"""
        print("Sending ENHANCED request to Gemini API for Overview & Comparison...")
        request_options = {"timeout": 300}
        response = model.generate_content(prompt, request_options=request_options)

        # --- Process Response ---
        if response.parts:
            generated_text = response.text
            print("Received response from Gemini.")

            # --- Parse the response into two sections ---
            overview_marker = "Topic Overview:"
            comparison_marker = "Detailed Comparative Analysis:"
            overview_start_index = generated_text.find(overview_marker)
            comparison_start_index = generated_text.find(comparison_marker)

            if overview_start_index != -1 and comparison_start_index != -1:
                topic_overview_result = generated_text[overview_start_index + len(overview_marker):comparison_start_index].strip()
                comparison_result = generated_text[comparison_start_index + len(comparison_marker):].strip()
                print("Successfully parsed Overview and Comparison sections.")
            elif overview_start_index != -1: # Found overview but not comparison
                topic_overview_result = generated_text[overview_start_index + len(overview_marker):].strip()
                if len(all_analyses) > 1: comparison_result = "Could not parse detailed comparison section."
                print("Parsed Overview section, Comparison marker not found.")
            else:
                print("Warning: Could not find expected markers. Using full text as overview.")
                topic_overview_result = generated_text
                comparison_result = "Parsing failed."

            # Reset comparison text if only one paper
            if len(all_analyses) == 1:
                comparison_result = "Comparison skipped: Only one paper was successfully analyzed."

        elif response.prompt_feedback:
             error_msg = f"Gemini request blocked ({response.prompt_feedback.block_reason})"
             print(error_msg)
             topic_overview_result = error_msg; comparison_result = error_msg
        else:
             print("Received empty response from Gemini.")
             topic_overview_result = "Overview failed: Empty response"; comparison_result = "Comparison failed: Empty response"

    except Exception as e:
        error_msg = f"Comparison/Overview generation failed due to error: {e}"
        print(f"ERROR during Gemini analysis: {type(e).__name__} - {e}")
        topic_overview_result = error_msg; comparison_result = error_msg

    return {"topic_overview": topic_overview_result, "comparison_result": comparison_result, "error": final_error}

# --- Report Generator Agent (Corrected Bold Formatting) ---
# --- (Ensure docx imports are present at the top) ---
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
import re # Import regular expressions for parsing points

# --- Helper function to add text as paragraphs/bullets ---
def add_text_with_bullets(document, text_block, style='List Bullet'):
    """Adds a block of text to the document, attempting to format lines
       starting with list markers (*, -, 1.) as bullets."""
    if not text_block:
        document.add_paragraph("N/A")
        return

    lines = text_block.strip().split('\n')
    for line in lines:
        stripped_line = line.strip()
        # Basic check for common bullet point/numbering patterns
        if re.match(r'^(\*|-|\d+\.\s)', stripped_line):
            # Remove the marker before adding as bullet
            content = re.sub(r'^(\*|-|\d+\.\s)\s*', '', stripped_line)
            document.add_paragraph(content, style=style)
        elif stripped_line: # Avoid adding empty paragraphs
            # Add as a normal paragraph if it's not a list item
            document.add_paragraph(stripped_line, style='Normal')
        # Add a check for empty lines if spacing is desired between points/paragraphs
        # elif not stripped_line and lines.index(line) > 0 and lines[lines.index(line)-1].strip():
        #      document.add_paragraph() # Add space after a content line

# --- Report Generator Agent (UPDATED for Points and URLs) ---
def generate_report_node(state: Dict):
    """
    Node function for the Report Generator Agent with point formatting
    and source URLs.
    """
    print("\n--- Node: Report Generator Agent (Points & URLs Mode) ---")
    topic = state.get("topic", "N/A")
    all_analyses = state.get("all_analyses", []) # Now contains pdf_url
    topic_overview = state.get("topic_overview", "Topic overview could not be generated.")
    comparison = state.get("comparison_result", "Detailed comparison could not be generated.")
    report_path = None
    final_error = state.get("error")

    if not all_analyses and not final_error:
        final_error = "Cannot generate report without successful summaries"
        return {"report_path": None, "error": final_error}
    elif not all_analyses and final_error:
        print("Skipping report generation due to prior error.")
        return {"report_path": None, "error": final_error}

    print(f"Generating formatted report for topic: '{topic}'")
    try:
        document = Document()
        # Set margins
        sections = document.sections
        for section in sections:
            section.left_margin = Inches(1.0); section.right_margin = Inches(1.0)
            section.top_margin = Inches(1.0); section.bottom_margin = Inches(1.0)

        # Set default font
        style = document.styles['Normal']; font = style.font
        font.name = 'Calibri'; font.size = Pt(11)

        # Define bullet style if needed (optional refinement)
        # bullet_style = document.styles.add_style('MyBullet', WD_STYLE_TYPE.PARAGRAPH)
        # bullet_style.base_style = document.styles['List Bullet']
        # bullet_font = bullet_style.font; bullet_font.name = 'Calibri'; bullet_font.size = Pt(11)
        # Use 'List Bullet' or your custom style name in add_text_with_bullets

        # --- Title ---
        title_h = document.add_heading(f"AI-Generated Research Report", 0)
        title_h.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        sub_title_h = document.add_heading(f"Topic: {topic}", 1)
        sub_title_h.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        document.add_paragraph()

        # --- Topic Overview Section ---
        document.add_heading("Topic Overview", level=2)
        add_text_with_bullets(document, topic_overview) # Use helper function
        document.add_page_break()

        # --- Comparative Analysis Section ---
        document.add_heading("Comparative Analysis", level=2)
        add_text_with_bullets(document, comparison) # Use helper function

        # --- Individual Summaries Section ---
        if all_analyses:
            document.add_page_break()
            document.add_heading("Individual Paper Summaries", level=2)
            for i, analysis in enumerate(all_analyses):
                document.add_heading(f"Paper {i+1}: {analysis.get('title', 'N/A')}", level=3)

                # Add Summary
                p_summary_label = document.add_paragraph(); p_summary_label.add_run("Summary:").bold = True
                add_text_with_bullets(document, analysis.get('summary', 'N/A'))

                # Add Methodology
                p_method_label = document.add_paragraph(); p_method_label.add_run("Methodology:").bold = True
                add_text_with_bullets(document, analysis.get('methodology', 'N/A'))

                # Add Key Findings
                p_findings_label = document.add_paragraph(); p_findings_label.add_run("Key Findings:").bold = True
                add_text_with_bullets(document, analysis.get('key_findings', 'N/A'))

                # Add space between papers only if not the last one
                if i < len(all_analyses) - 1: document.add_paragraph()
        else:
             document.add_paragraph("\nNo individual paper summaries could be generated.")


        # --- ADD SOURCE URLS SECTION (NEW) ---
        document.add_page_break()
        document.add_heading("Source Paper URLs", level=2)
        if all_analyses:
            for i, analysis in enumerate(all_analyses):
                title = analysis.get('title', f'Paper {i+1}')
                url = analysis.get('pdf_url', 'URL not available')
                p = document.add_paragraph(style='List Bullet')
                p.add_run(f"{title}: ").italic = True
                # Add URL (basic text, not hyperlink for simplicity now)
                p.add_run(url)
        else:
            document.add_paragraph("No source papers were processed.")


        # --- Save ---
        report_path = os.path.join(os.getcwd(), OUTPUT_REPORT_NAME)
        document.save(report_path)
        print(f"Formatted report with points and URLs saved to: '{report_path}'")

    except Exception as e:
        print(f"ERROR during DOCX generation: {type(e).__name__} - {e}")
        final_error = final_error or f"Report generation failed: {e}"
        return {"report_path": None, "error": final_error}

    return {"report_path": report_path, "error": final_error}

# --- Build the Graph (No changes needed) ---
print("\n--- Building LangGraph Workflow ---")
workflow = StateGraph(ResearchState)
workflow.add_node("search_papers", search_academic_papers_node)
workflow.add_node("download_parse", download_and_parse_pdfs_node)
workflow.add_node("summarize_papers", summarize_papers_node)
workflow.add_node("compare_analyses", compare_analyses_node)
workflow.add_node("generate_report", generate_report_node)
workflow.set_entry_point("search_papers")
workflow.add_edge("search_papers", "download_parse")
workflow.add_edge("download_parse", "summarize_papers")
workflow.add_edge("summarize_papers", "compare_analyses")
workflow.add_edge("compare_analyses", "generate_report")
workflow.add_edge("generate_report", END)
app = workflow.compile()
print("--- Workflow Compiled ---")

# --- Run the Graph (No changes needed) ---
if __name__ == "__main__":
    print("\n--- Starting Research Workflow ---")
    initial_topic = "Agentic AI workflows using LangGraph"
    inputs = {"topic": initial_topic, "paper_limit": 3}

    final_state = None
    try:
        print("Streaming graph execution events...")
        for event in app.stream(inputs):
            for node_name, output_state_update in event.items():
                print(f"--- Finished Node: {node_name} ---")
                if isinstance(output_state_update, dict):
                    updated_keys = list(output_state_update.keys())
                    print(f"  State keys updated: {updated_keys}")
                    if "error" in output_state_update and output_state_update["error"]:
                        print(f"  !! ERROR RECORDED by node '{node_name}': {output_state_update['error']}")
                if node_name == "__end__":
                     final_state = output_state_update
                     print("  Workflow has reached the END state.")
    except Exception as e:
        print(f"\n--- Workflow Execution Error ---")
        print(f"Critical Error during graph execution: {e}")

    # --- Print Final Results (No changes needed) ---
    print("\n--- Workflow Execution Finished ---")
    if final_state:
        print("\nFinal State Summary:")
        print(f"- Topic: {final_state.get('topic')}")
        print(f"- Report Path: {final_state.get('report_path', 'Not Generated')}")
        if final_state.get('error'):
             print(f"- Last Recorded Error: {final_state.get('error')}")
        elif not final_state.get('report_path'):
             print("- Status: Completed Successfully.")
        else:
             print("- Status: Completed Successfully.")
    else:
        print("WStatus: Completed Successfully.")
    print("---------------------------------")