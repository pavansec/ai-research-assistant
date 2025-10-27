from docx import Document
from docx.shared import Inches, Pt # For setting font size, margins etc.
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT # For text alignment
import os

# --- Define the Report Generator Agent Function ---
def create_report_docx(topic: str, analysis_list: list, comparison: str, output_filename: str = "AI_Generated_Research_Report.docx"):
    """
    Creates a Word document (.docx) report from paper analyses and comparison.

    Args:
        topic: The original research topic searched.
        analysis_list: A list of dictionaries (from SummarizerAgent).
                       Each dict has 'summary', 'methodology', 'key_findings'.
        comparison: The comparative analysis string (from ComparatorAgent).
        output_filename: The name for the output .docx file.

    Returns:
        The path to the generated docx file, or None if an error occurs.
    """
    print(f"\n--- Running Report Generator Agent ---")
    print(f"Generating report for topic: '{topic}'")

    if not analysis_list:
        print("ERROR: No paper analyses provided to generate report.")
        print("--- Report Generator Failed ---")
        return None
    if not comparison:
        print("ERROR: No comparative analysis provided.")
        print("--- Report Generator Failed ---")
        return None

    try:
        # --- Create a new Document ---
        document = Document()

        # --- Set basic styles (optional) ---
        style = document.styles['Normal']
        font = style.font
        font.name = 'Calibri' # Example font
        font.size = Pt(11)

        # --- Add Title ---
        title = document.add_heading(f"AI-Generated Research Report", level=0) # Main title
        title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        document.add_heading(f"Topic: {topic}", level=1) # Subtitle with topic
        document.add_paragraph() # Add some space

        # --- Add Comparative Analysis Section ---
        document.add_heading("Comparative Analysis", level=2)
        document.add_paragraph(comparison)
        document.add_page_break() # Start individual papers on a new page

        # --- Add Individual Paper Analysis Sections ---
        document.add_heading("Individual Paper Summaries", level=2)
        for i, analysis in enumerate(analysis_list):
            document.add_heading(f"Paper {i+1} Analysis", level=3)

            # Add Summary
            summary_p = document.add_paragraph()
            summary_p.add_run("Summary:").bold = True
            summary_p.add_run(f"\n{analysis.get('summary', 'N/A')}") # Add newline for clarity

            # Add Methodology
            methodology_p = document.add_paragraph()
            methodology_p.add_run("Methodology:").bold = True
            methodology_p.add_run(f"\n{analysis.get('methodology', 'N/A')}")

            # Add Key Findings
            findings_p = document.add_paragraph()
            findings_p.add_run("Key Findings:").bold = True
            findings_p.add_run(f"\n{analysis.get('key_findings', 'N/A')}")

            document.add_paragraph() # Add space between papers

        # --- Save the Document ---
        save_path = os.path.join(os.getcwd(), output_filename) # Save in current dir
        document.save(save_path)
        print(f"Report successfully saved to: '{save_path}'")

    except Exception as e:
        print(f"\n--- Error during DOCX generation ---")
        print(f"Error details: {type(e).__name__} - {e}")
        print("Troubleshooting tips:")
        print(" - Ensure 'python-docx' library is installed correctly.")
        print(" - Check file permissions in the output directory.")
        print("-----------------------------------")
        return None # Return None on error

    print("--- Report Generator Agent Complete ---")
    return save_path

# --- Run a Test ---
if __name__ == "__main__":
    # --- Use the same sample data from Comparator Agent ---
    test_topic_main = "Agentic AI workflows using LangGraph" # Original Topic

    analysis_1_main = {
        'summary': """This paper proposes a real-time Sign Language to Speech Conversion system specifically designed for Tamil-speaking communities. The system captures hand gestures via a standard webcam, utilizes MediaPipe for accurate hand landmark detection, and then classifies these gestures using a Convolutional Neural Network (CNN) trained on Tamil sign language. The recognized Tamil alphabets or words are subsequently converted into spoken Tamil using the Google Text-to-Speech (gTTS) API. This accessible and cost-effective solution aims to overcome communication barriers for individuals with hearing or speech impairments in Tamil-speaking regions.""",
        'methodology': """The primary methodology involves a multi-stage machine learning and computer vision approach. It begins with data collection of Tamil sign language gestures using a webcam. Preprocessing involves extracting 21 hand landmarks using MediaPipe and normalizing them. These landmarks are then fed into a Convolutional Neural Network (CNN) for model training and real-time gesture recognition, classifying the gestures into specific Tamil alphabets or words. Finally, the recognized text is converted into spoken Tamil through the Speech Synthesis Module utilizing the gTTS API, enabling real-time audio output.""",
        'key_findings': """1. The system achieved a high accuracy of 92% in controlled settings and 85% in real-world environments for Tamil sign language gesture recognition. 2. It consistently provided speech output with minimal latency, typically within 1-2 seconds after gesture recognition, making it practical for real-time communication."""
    }

    analysis_2_main = {
        'summary': """This study explores the use of Long Short-Term Memory (LSTM) networks combined with computer vision for recognizing dynamic American Sign Language (ASL) gestures from video sequences. It focuses on capturing temporal dependencies in sign movements. The system uses OpenPose for body and hand keypoint extraction and feeds these sequences into an LSTM model for classification. Text-to-speech is handled by a standard offline library.""",
        'methodology': """The methodology uses video input, processed frame-by-frame with OpenPose to extract skeletal keypoints. These time-series keypoint sequences are then fed into a Long Short-Term Memory (LSTM) recurrent neural network trained to classify dynamic ASL gestures. Speech synthesis is done using an offline text-to-speech engine.""",
        'key_findings': """1. The LSTM model achieved 88% accuracy on a dataset of dynamic ASL signs. 2. Capturing temporal information via LSTM significantly outperformed static image classifiers for dynamic gestures. 3. OpenPose keypoints provided a robust feature set, reducing sensitivity to background variations."""
    }

    # IMPORTANT: Copy the full comparison text from your previous run here!
    comparison_text_main = """
Here's a comparative analysis of the two research paper summaries:

---

### Comparative Analysis

These two papers both explore the domain of sign language recognition and conversion to speech, leveraging computer vision and machine learning techniques to bridge communication gaps.

**1. Similarities:**

* **Core Goal:** Both aim to facilitate communication for individuals with hearing/speech impairments by converting sign language gestures into spoken output.
* **Input Modality:** Both utilize visual input (webcam/video sequences) for gesture capture.
* **Feature Extraction:** Both employ sophisticated computer vision libraries for extracting skeletal/hand keypoints/landmarks from the visual input (MediaPipe in Paper 1, OpenPose in Paper 2).
* **Machine Learning for Classification:** Both rely on deep learning models (CNN in Paper 1, LSTM in Paper 2) for classifying recognized gestures.
* **Output Stage:** Both incorporate a Text-to-Speech (TTS) module to convert the recognized signs (as text) into spoken language.
* **Application Area:** Both contribute to assistive technologies and human-computer interaction for the deaf and hard-of-hearing community.

**2. Differences:**

* **Target Sign Language & Gesture Type:**
    * **Paper 1:** Focuses on **Tamil Sign Language**, recognizing primarily individual **alphabets or words**, implying a potential focus on more static or discrete signs.
    * **Paper 2:** Concentrates on **American Sign Language (ASL)** and explicitly addresses **dynamic gestures** from video sequences, highlighting the importance of temporal dependencies in sign movements.
* **Machine Learning Model:**
    * **Paper 1:** Uses a **Convolutional Neural Network (CNN)**, which is typically well-suited for spatial feature extraction from images, but can be adapted for sequences.
    * **Paper 2:** Employs a **Long Short-Term Memory (LSTM) network**, which is specifically designed for sequential data, making it highly appropriate for capturing the temporal information inherent in dynamic gestures.
* **Keypoint Extraction Tool:**
    * **Paper 1:** Uses **MediaPipe**, primarily for hand landmark detection.
    * **Paper 2:** Uses **OpenPose**, which extracts both body and hand keypoints, suggesting it might capture a broader range of motion and context crucial for dynamic ASL signs.
* **Text-to-Speech Implementation:**
    * **Paper 1:** Utilizes the **Google Text-to-Speech (gTTS) API**, implying an online, real-time integration for speech synthesis.
    * **Paper 2:** Mentions a "standard **offline** library," suggesting less emphasis on the real-time nature or specific technology of the TTS component.
* **Performance Emphasis & Context:**
    * **Paper 1:** Highlights performance in both *controlled* (92%) and *real-world* (85%) environments for Tamil Sign Language, along with minimal latency (1-2 seconds), emphasizing practical, real-time application.
    * **Paper 2:** Reports 88% accuracy for dynamic ASL and stresses that capturing *temporal information via LSTM significantly outperformed static image classifiers*, and that OpenPose provided a robust feature set, focusing on the methodological advantage for dynamic signs.

**3. Overall Theme/Contribution:**

Both papers make significant contributions to the field of sign language recognition, but with different primary emphases.

* **Paper 1's** main contribution is providing an **accessible, cost-effective, and real-time solution** for **Tamil Sign Language to Speech conversion**, specifically tailored for a regional community. Its strength lies in demonstrating practical implementation and real-world utility with measurable latency.
* **Paper 2's** main contribution is advancing the understanding and methodology for recognizing **dynamic sign language gestures** by effectively leveraging **LSTM networks for temporal modeling** and demonstrating the efficacy of *OpenPose* for robust feature extraction in this complex task. It focuses on tackling the inherent challenges of movement and sequence in sign language.

In essence, Paper 1 delivers a practical, region-specific application, while Paper 2 focuses on a methodological advancement for a more complex type of sign language recognition.
    """

    test_analyses_main = [analysis_1_main, analysis_2_main]

    # --- Generate the report ---
    generated_file_path = create_report_docx(test_topic_main, test_analyses_main, comparison_text_main)

    if generated_file_path:
        print(f"\nReport generation test successful. File saved at: {generated_file_path}")
    else:
        print("\nReport generation test failed.")